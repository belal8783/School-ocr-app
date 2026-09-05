import os
import io
import re
import time
import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from PIL import Image, ImageOps
from pdf2image import convert_from_bytes

# ---------------------------------------------------------
# Unicode Bengali Cleaning Routine
# ---------------------------------------------------------
def clean_bengali_symbols(text):
    if not text:
        return ""
    cleaned_text = re.sub(r'\u25cc', '', text)
    return cleaned_text.replace('◌', '')

def is_arabic(text):
    return bool(re.search(r'[\u0600-\u06FF]', text))

# ---------------------------------------------------------
# Page Config & Branding
# ---------------------------------------------------------
st.set_page_config(
    page_title="Handwritten to Word & PDF Agent",
    page_icon="📝",
    layout="wide"
)

st.title("📝 School Sheet Handwritten to Word & PDF Agent")
st.caption("Developed by: Belal Hossain | Gemini 3.6 Flash Powered")

# ---------------------------------------------------------
# Sidebar Options
# ---------------------------------------------------------
st.sidebar.header("⚙️ কাস্টমাইজেশন সেটিং")

font_size = st.sidebar.slider("মূল টেক্সট ফন্ট সাইজ (Font Size)", min_value=10, max_value=24, value=13, step=1)

bangla_font_name = st.sidebar.selectbox(
    "বাংলা ফন্ট নির্বাচন করুন (ওয়ার্ডের জন্য)",
    ["Kalpurush", "SolaimanLipi", "Siyam Rupali", "Arial"]
)

english_font_name = st.sidebar.selectbox(
    "ইংরেজি ফন্ট নির্বাচন করুন",
    ["Times New Roman", "Calibri", "Arial"]
)

arabic_font_name = st.sidebar.selectbox(
    "আরবি ফন্ট নির্বাচন করুন",
    ["Traditional Arabic", "Amiri", "Scheherazade"]
)

# ---------------------------------------------------------
# API Key Configuration
# ---------------------------------------------------------
api_key = st.secrets.get("GEMINI_API_KEY") or os.getenv("GEMINI_API_KEY")

if not api_key:
    api_key = st.sidebar.text_input("Gemini API Key দিন:", type="password")

if not api_key:
    st.info("⚠️ অ্যাপটি চালাতে Gemini API Key প্রয়োজন।")
    st.stop()

genai.configure(api_key=api_key)

# ---------------------------------------------------------
# Model Selection (FIXED: gemini-3.6-flash)
# ---------------------------------------------------------
MODEL_NAME = 'gemini-3.6-flash'

model = genai.GenerativeModel(
    model_name=MODEL_NAME,
    generation_config={
        "temperature": 0.1,
        "top_p": 0.95,
        "max_output_tokens": 4096,
    }
)

# ---------------------------------------------------------
# Safety Settings
# ---------------------------------------------------------
safety_settings = [
    {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
    {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
    {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
    {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"}
]

# ---------------------------------------------------------
# File Upload & Session State Setup
# ---------------------------------------------------------
uploaded_files = st.file_uploader(
    "আপনার শিটের ছবি (JPG, PNG) অথবা PDF ফাইল আপলোড করুন",
    type=["jpg", "jpeg", "png", "pdf"],
    accept_multiple_files=True
)

custom_filename = st.text_input("ডাউনলোড ফাইলের নাম লিখুন:", value="Converted_School_Sheet")

if 'final_converted_text' not in st.session_state:
    st.session_state.final_converted_text = ""
if 'conversion_done' not in st.session_state:
    st.session_state.conversion_done = False

# ---------------------------------------------------------
# Master AI Prompt
# ---------------------------------------------------------
SYSTEM_PROMPT = """
আপনি একজন বিশেষজ্ঞ বাংলা OCR এবং ট্রান্সক্রিপশন সিস্টেম।
ছবিতে থাকা বাংলা এবং ইংরেজি হাতে লেখা ও ছাপানো টেক্সট নিখুঁতভাবে রূপান্তর করুন।

বিশেষ নির্দেশাবলী:
১. ছবিতে যদি কোনো ছক বা টেবিল থাকে (যেমন: প্রদত্ত শব্দ | শব্দের অর্থ), তাহলে সেটিকে অবশ্যই স্ট্যান্ডার্ড Markdown Table আকারে লিখবেন।
উদাহরণ:
| প্রদত্ত শব্দ | শব্দের অর্থ |
| --- | --- |
| অসি | কলম |
| ইতি | শেষ |

২. ছবিতে কোনো লাল কালির বা কাটাকাটি লেখা থাকলে, কাটাকাটির ভেতরের সঠিক শব্দটি পড়ে সুন্দরভাবে টাইপ করবেন।
৩. অতিরিক্ত কোনো কথা বা ব্যাখ্যা দেবেন না। শুধুমাত্র ছবিতে থাকা পড়া অংশটি সঠিকভাবে সাজিয়ে আউটপুট দিন।
"""

# ---------------------------------------------------------
# Image Compressor and Optimizer
# ---------------------------------------------------------
def optimize_and_convert_image(pil_img):
    # অটো ওরিয়েন্টেশন ফিক্স
    pil_img = ImageOps.exif_transpose(pil_img)
    
    # সাইজ অপটিমাইজ করা (Max 1280px)
    max_size = (1280, 1280)
    pil_img.thumbnail(max_size, Image.Resampling.LANCZOS)
    
    img_byte_arr = io.BytesIO()
    pil_img.convert('RGB').save(img_byte_arr, format='JPEG', quality=85)
    
    return {
        'mime_type': 'image/jpeg',
        'data': img_byte_arr.getvalue()
    }

# ---------------------------------------------------------
# Word Document Generator with Grid Borders
# ---------------------------------------------------------
def create_word_docx(text, bangla_font, english_font, arabic_font, font_size):
    doc = Document()
    lines = text.split('\n')
    
    in_table = False
    table_data = []

    def flush_table(t_data):
        if not t_data:
            return
        
        valid_rows = [r for r in t_data if not all(re.match(r'^[\s:-]+$', cell) for cell in r)]
        if not valid_rows:
            return
            
        max_cols = max(len(r) for r in valid_rows)
        word_table = doc.add_table(rows=len(valid_rows), cols=max_cols)
        word_table.style = 'Table Grid'
        word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for r_idx, row in enumerate(valid_rows):
            for c_idx, cell_value in enumerate(row):
                if c_idx < max_cols:
                    cell = word_table.cell(r_idx, c_idx)
                    cell.text = cell_value.strip()
                    
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.font.size = Pt(font_size)
                            if is_arabic(cell_value):
                                run.font.name = arabic_font
                            elif cell_value.isascii():
                                run.font.name = english_font
                            else:
                                run.font.name = bangla_font
                                
        doc.add_paragraph()

    for line in lines:
        stripped = line.strip()
        
        if stripped.startswith('|') and stripped.endswith('|'):
            in_table = True
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            table_data.append(cells)
            continue
        else:
            if in_table:
                flush_table(table_data)
                table_data = []
                in_table = False

        if not stripped:
            continue
            
        p = doc.add_paragraph()
        is_heading = line.startswith('#')
        current_line = re.sub(r'^#+\s*', '', line) if is_heading else line
            
        parts = re.split(r'(<u>.*?</u>|\[ইমেজ নোট:.*?\])', current_line)
        
        for part in parts:
            if not part:
                continue
                
            run = p.add_run()
            run.text = part
            run.font.bold = is_heading
            run.font.size = Pt(font_size + 3) if is_heading else Pt(font_size)
            
            if is_arabic(part):
                run.font.name = arabic_font
            elif part.isascii():
                run.font.name = english_font
            else:
                run.font.name = bangla_font

    if in_table:
        flush_table(table_data)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ---------------------------------------------------------
# PDF Document Generator
# ---------------------------------------------------------
def create_pdf(text, font_size):
    buffer = io.BytesIO()
    pdf = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    
    custom_style = ParagraphStyle(
        'CustomStyle',
        parent=styles['Normal'],
        fontSize=font_size,
        leading=font_size + 5
    )
    
    story = []
    for line in text.split('\n'):
        if line.strip():
            safe_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(safe_line, custom_style))
            story.append(Spacer(1, 4))
            
    pdf.build(story)
    buffer.seek(0)
    return buffer

# ---------------------------------------------------------
# Execution Logic
# ---------------------------------------------------------
if uploaded_files and st.button("🚀 কনভার্ট শুরু করুন"):
    combined_result = ""
    images_to_process = []

    with st.spinner("ফাইল মেমোরিতে প্রসেস করা হচ্ছে..."):
        for uploaded_file in uploaded_files:
            if uploaded_file.name.lower().endswith(".pdf"):
                pdf_bytes = uploaded_file.read()
                converted_images = convert_from_bytes(pdf_bytes)
                images_to_process.extend(converted_images)
            else:
                uploaded_file.seek(0)
                img = Image.open(uploaded_file)
                images_to_process.append(img)

    total_pages = len(images_to_process)
    progress_bar = st.progress(0)
    status_text = st.empty()

    for index, img in enumerate(images_to_process):
        success = False
        optimized_payload = optimize_and_convert_image(img)
        
        for attempt in range(3):
            status_text.text(f"প্রসেসিং চলছে: পৃষ্ঠা {index + 1} / {total_pages} (ট্রাই: {attempt + 1})")
            try:
                response = model.generate_content(
                    [SYSTEM_PROMPT, optimized_payload],
                    safety_settings=safety_settings
                )
                
                if response and hasattr(response, 'text') and response.text.strip():
                    page_text = clean_bengali_symbols(response.text)
                    combined_result += f"\n\n--- পৃষ্ঠা {index + 1} ---\n\n" + page_text
                    success = True
                    break
                else:
                    time.sleep(2)
                
            except Exception as e:
                time.sleep(3)
        
        if not success:
            combined_result += f"\n\n--- পৃষ্ঠা {index + 1} ---\n\n[এরর: পৃষ্ঠাটি প্রসেস করা সম্ভব হয়নি।]"
            
        progress_bar.progress((index + 1) / total_pages)

    status_text.empty()
    
    st.session_state.final_converted_text = combined_result
    st.session_state.conversion_done = True

# ---------------------------------------------------------
# Display Output & Downloads
# ---------------------------------------------------------
if st.session_state.conversion_done:
    st.success("✅ কনভার্ট সফলভাবে সম্পন্ন হয়েছে!")

    st.subheader("📝 কনভার্ট হওয়া টেক্সট প্রিভিউ (ছকসহ):")
    st.markdown(st.session_state.final_converted_text)

    col1, col2 = st.columns(2)
    
    word_file = create_word_docx(st.session_state.final_converted_text, bangla_font_name, english_font_name, arabic_font_name, font_size)
    pdf_file = create_pdf(st.session_state.final_converted_text, font_size)

    with col1:
        st.download_button(
            label="📄 Word (.docx) ডাউনলোড করুন",
            data=word_file,
            file_name=f"{custom_filename}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    with col2:
        st.download_button(
            label="📕 PDF (.pdf) ডাউনলোড করুন",
            data=pdf_file,
            file_name=f"{custom_filename}.pdf",
            mime="application/pdf"
        )
